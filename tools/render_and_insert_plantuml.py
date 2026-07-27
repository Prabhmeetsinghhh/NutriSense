from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Inches


def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def insert_images(input_docx: Path, output_docx: Path, diagram_dir: Path) -> None:
    doc = Document(str(input_docx))

    mapping = [
        ("2.2 Use Case diagram", ["01_use_case_diagram.png"]),
        ("2.3 ER Model", ["04_er_diagram.png"]),
        ("2.4 Data flow Diagram", ["02_dfd_level_0_context.png", "03_dfd_level_1.png"]),
        ("2.5 Control Flow Diagram", ["09_control_flow_diagram.png"]),
        ("2.6 State Transition Diagram", ["08_state_transition_diagram.png"]),
        ("3.1.1 System Architectural Diagram", ["13_deployment_diagram.png"]),
        ("3.3.1 Flowchart", ["10_flowchart_component_logic.png"]),
        ("2.3 Activity Diagram", ["05_activity_diagram.png"]),
        ("2.4 Sequence Diagram", ["06_sequence_diagram.png"]),
        ("2.5 Class Diagram", ["07_class_diagram.png"]),
        ("3.3.1 Package Diagram", ["11_package_diagram.png"]),
        ("3.3.2 Component Diagram", ["12_component_diagram.png"]),
        ("3.3.2 Deployment Diagram", ["13_deployment_diagram.png"]),
    ]

    inserted = set()
    for para in doc.paragraphs:
        text = para.text.strip()
        for heading, images in mapping:
            if heading in text and heading not in inserted:
                current_anchor = para
                for idx, image_name in enumerate(images, start=1):
                    image_path = diagram_dir / image_name
                    if image_path.exists():
                        caption = f"Figure {heading}{' (Part ' + str(idx) + ')' if len(images) > 1 else ''}"
                        cap_para = insert_paragraph_after(current_anchor, caption)
                        img_para = insert_paragraph_after(cap_para)
                        run = img_para.add_run()
                        run.add_picture(str(image_path), width=Inches(6.7))
                        current_anchor = img_para
                inserted.add(heading)
                break

    doc.save(str(output_docx))


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    input_docx = root / "NutriSense_SDS_Modified.docx"
    output_docx = root / "NutriSense_SDS_With_Correct_Diagrams.docx"
    rendered_dir = root / "docs" / "uml-pro" / "rendered"

    if not input_docx.exists():
        raise FileNotFoundError(f"Input not found: {input_docx}")
    if not rendered_dir.exists():
        raise FileNotFoundError(f"Rendered diagram folder not found: {rendered_dir}")

    insert_images(input_docx, output_docx, rendered_dir)
    print(f"Created: {output_docx}")


if __name__ == "__main__":
    main()
