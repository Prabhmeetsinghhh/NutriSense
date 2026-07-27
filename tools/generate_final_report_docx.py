from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from datetime import date

BASE_TEMPLATE = r"C:\Users\DELL\Desktop\MinorProject\BTech_CSE_MINI_Project_Report_Format.docx"
OUTPUT_FILE = r"C:\Users\DELL\Desktop\MinorProject\NutriSense_Final_Mini_Project_Report.docx"

TITLE = "NUTRISENSE: PERSONALIZED DIET AND FITNESS PLANNER FOR INDIAN USERS"
STUDENT_NAME = "Prabhmeet Singh"
ROLL_NO = "EN23CS301742"
GUIDES = "Vishal Sharma Sir and Anusha Jain Mam"
DEPARTMENT = "Computer Science & Engineering"
MONTH_YEAR = "APRIL 2026"


# Report sections authored from the implemented codebase and synopsis.
SECTIONS = [
    (
        "Abstract",
        [
            "NutriSense is a full-stack web application that generates personalized diet and fitness plans for Indian users based on age, body metrics, activity profile, goal, and diet preference. The system addresses practical gaps in generic planning apps by combining Indian meal options, budget-aware suggestions, and beginner-friendly workout scheduling.",
            "The backend is implemented with FastAPI and rule-based services for BMI estimation, goal mapping, calorie targeting, protein computation, and meal distribution across breakfast, lunch, evening meal, and dinner. A weekly fitness schedule is generated in parallel to maintain a balanced health plan. The frontend is built with React, TypeScript, and Vite, providing form-based input capture and structured visualization of generated plans.",
            "The project demonstrates a working end-to-end architecture with validated API integration, clean user flow, and actionable recommendations. The output includes macros, per-meal guidance, estimated costs, and structured workout routines. NutriSense is designed as an affordable and practical solution for students and working professionals who need localized, implementable health planning.",
            "Keywords: personalized nutrition, Indian diet planning, rule-based recommendation, BMI, TDEE, FastAPI, React, fitness scheduling"
        ],
    ),
    (
        "Chapter 1: Introduction",
        [
            "1.1 Introduction",
            "Health planning applications are common, but many are not adapted to Indian food patterns, student budgets, and day-to-day lifestyle constraints. NutriSense solves this by generating realistic meal and workout plans from user-specific inputs such as age, height, weight, fitness level, and diet preference.",
            "1.2 Literature Review",
            "Prior work in personalized nutrition consistently uses BMI and energy expenditure methods to estimate dietary requirements. Rule-based systems remain effective where explainability and predictable behavior are required. Existing commercial applications provide broad recommendations but often lack localization for Indian meals and affordability bands.",
            "1.3 Objectives",
            "The objective is to design and implement a practical planner that (a) computes BMI, calorie and protein targets, (b) generates four daily meals with macros and cost hints, (c) supports veg, non-veg, vegan, and eggetarian preferences, and (d) creates a weekly fitness routine aligned to user goals.",
            "1.4 Significance",
            "NutriSense provides an integrated and transparent planning workflow. It improves accessibility by presenting familiar foods, simple home/gym exercises, and budget-aware recommendations.",
            "1.5 Research Design",
            "The project follows an applied design-and-build approach with iterative validation. Rule sets were refined against expected outputs and integrated with a responsive frontend.",
            "1.6 Source of Data",
            "Meal entries, macro approximations, and cost ranges were curated in structured Python dictionaries for Indian foods. User data is collected through frontend forms and processed through backend services.",
            "1.7 Chapter Scheme",
            "Chapter 1 introduces the problem. Chapter 2 defines requirements. Chapter 3 presents system design. Chapter 4 discusses implementation and testing. Chapter 5 shows results. Chapter 6 gives conclusions. Chapter 7 describes future scope.",
        ],
    ),
    (
        "Chapter 2: Requirements Specification",
        [
            "2.1 User Characteristics",
            "Target users include college students, hostel residents, and working professionals seeking practical health planning with limited time and budget.",
            "2.2 Functional Requirements",
            "The system shall collect user profile data, compute BMI and targets, generate a four-meal plan, produce a seven-day fitness plan, and render details in a user-friendly dashboard.",
            "2.3 Dependencies",
            "Frontend depends on React, TypeScript, Vite, axios, and route management. Backend depends on FastAPI and Uvicorn. API communication is JSON over HTTP.",
            "2.4 Performance Requirements",
            "Plan generation should complete within interactive response times for normal local development usage and should provide deterministic outputs for equivalent inputs.",
            "2.5 Hardware Requirements",
            "Minimum requirement: modern laptop/desktop with 8 GB RAM, Python 3.11 environment for backend, Node.js runtime for frontend, and browser support for local deployment.",
            "2.6 Constraints and Assumptions",
            "Recommendations are general wellness suggestions and not a substitute for medical diagnosis. Macro and cost values are approximate and designed for practical guidance.",
        ],
    ),
    (
        "Chapter 3: System Design",
        [
            "3.1 Algorithmic Basis",
            "Diet planning uses Mifflin-St Jeor based TDEE estimation, BMI-based goal mapping, protein multipliers by experience level, and fixed per-meal protein distribution.",
            "3.2 Functional Design",
            "Input processing, rule execution, meal assembly, fitness scheduling, and response formatting are modularized in backend services and router-level transformations.",
            "3.3 System Design",
            "3.3.1 Data Flow",
            "User enters profile details in frontend -> frontend sends payload to backend API -> diet and fitness services compute recommendations -> structured JSON is returned -> frontend renders meal cards and workout plan.",
            "3.3.2 Activity Flow",
            "Authenticate (basic flow), enter profile, choose goal and diet type, submit form, review personalized plan, and iterate with updated inputs.",
            "3.3.3 Class and Module View",
            "Key modules include IndianDietService, fitnessService functions, Indian foods model dictionaries, and apiRouter orchestration.",
            "3.4 Database Design",
            "This version uses in-code data structures instead of a relational database. The design is intentionally lightweight for fast iteration and deterministic testing.",
        ],
    ),
    (
        "Chapter 4: Implementation, Testing, and Maintenance",
        [
            "4.1 Languages, IDEs, Tools, and Technologies",
            "Backend: Python, FastAPI, Uvicorn. Frontend: React, TypeScript, Vite. Development: VS Code, npm scripts, virtual environment setup.",
            "4.2 Testing Techniques and Test Plan",
            "Manual API and UI tests were performed across user combinations for goal, fitness level, and diet type. Validation included output shape checks, meal coverage, and mapping consistency.",
            "4.3 Installation Instructions",
            "Backend: install requirements and run Uvicorn app.main:app on localhost:8000. Frontend: install dependencies and run Vite development server.",
            "4.4 End User Instructions",
            "Open the frontend page, fill required profile fields, select diet preference and goal, click generate plan, and inspect macros, cost estimates, and weekly exercise schedule.",
            "4.5 Maintenance Approach",
            "Future maintenance can add persistent storage, user accounts, stricter validation, and expanded localized food datasets while preserving current API contracts.",
        ],
    ),
    (
        "Chapter 5: Results and Discussion",
        [
            "5.1 User Interface Representation",
            "The interface presents a guided multi-page flow: landing, login, user details form, and plan result dashboard.",
            "5.2 Module-wise Results",
            "Diet module successfully computes BMI, calories, protein targets, and four-meal plans. Fitness module returns a weekly day-wise routine with intensity-aware suggestions.",
            "5.3 Sample Output Summary",
            "For representative users, the system produced complete breakfast, lunch, evening, and dinner recommendations with macro values and cost ranges, plus a seven-day exercise schedule.",
            "5.4 Backend Representation",
            "The API endpoint generates an aggregated response containing diet plan, fitness plan, user context, and explanatory metadata used by the frontend.",
            "5.5 Discussion",
            "The rule-based design improves transparency and consistency. Localization and budget tiering improve relevance over generic recommendations.",
        ],
    ),
    (
        "Chapter 6: Summary and Conclusions",
        [
            "NutriSense delivers an end-to-end, personalized, and localized planning experience for Indian users. The project integrates diet and fitness recommendations in one workflow with explainable rule logic and clean frontend interaction.",
            "The implementation meets core objectives: profile-driven recommendation, macro-aware four-meal generation, weekly workout scheduling, and budget-sensitive outputs. The system is usable in its current state and suitable as a base for advanced personalization.",
        ],
    ),
    (
        "Chapter 7: Future Scope",
        [
            "Add user authentication and persistent plan history.",
            "Integrate real nutrition databases and dynamic meal substitution.",
            "Use machine learning for preference adaptation while retaining rule safety constraints.",
            "Introduce wearable integration and progress analytics dashboards.",
            "Add multilingual UI and deeper regional food coverage across India.",
        ],
    ),
    (
        "Bibliography",
        [
            "[1] World Health Organization, Obesity and Overweight, Fact Sheet.",
            "[2] Mifflin, M. D. et al., A New Predictive Equation for Resting Energy Expenditure in Healthy Individuals, The American Journal of Clinical Nutrition, 1990.",
            "[3] National Institutes of Health, Clinical Guidelines on Identification, Evaluation, and Treatment of Overweight and Obesity in Adults.",
            "[4] Harvard T.H. Chan School of Public Health, The Nutrition Source.",
            "[5] Research literature on personalized nutrition and adherence outcomes.",
        ],
    ),
]


def remove_all_content(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        body.remove(child)


def add_center_paragraph(doc: Document, text: str, bold: bool = False, size: int = 12, space_after: int = 6) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def setup_default_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)


def add_heading_line(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def add_body_line(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)


def build_report() -> None:
    doc = Document(BASE_TEMPLATE)
    setup_default_style(doc)
    remove_all_content(doc)

    # Front cover page
    add_center_paragraph(doc, TITLE, bold=True, size=18, space_after=24)
    add_center_paragraph(doc, "A Minor Project Report", size=14)
    add_center_paragraph(doc, "Submitted in partial fulfillment of requirement of the", size=12)
    add_center_paragraph(doc, "Degree of", size=12)
    add_center_paragraph(doc, "BACHELOR OF TECHNOLOGY in COMPUTER SCIENCE & ENGINEERING", bold=True, size=13, space_after=14)
    add_center_paragraph(doc, "BY", size=12)
    add_center_paragraph(doc, STUDENT_NAME, bold=True, size=13)
    add_center_paragraph(doc, ROLL_NO, bold=True, size=12, space_after=14)
    add_center_paragraph(doc, "Under the Guidance of", size=12)
    add_center_paragraph(doc, GUIDES, bold=True, size=12, space_after=14)
    add_center_paragraph(doc, "Department of Computer Science & Engineering", size=12)
    add_center_paragraph(doc, "Faculty of Engineering", size=12)
    add_center_paragraph(doc, "MEDICAPS UNIVERSITY, INDORE - 453331", bold=True, size=12)
    add_center_paragraph(doc, MONTH_YEAR, bold=True, size=12)

    add_page_break(doc)

    # Title page
    add_center_paragraph(doc, TITLE, bold=True, size=16, space_after=16)
    add_center_paragraph(doc, "A Minor Project Report", size=13)
    add_center_paragraph(doc, "Submitted in partial fulfillment of requirement of the", size=12)
    add_center_paragraph(doc, "Degree of", size=12)
    add_center_paragraph(doc, "BACHELOR OF TECHNOLOGY in COMPUTER SCIENCE & ENGINEERING", bold=True, size=12, space_after=12)
    add_center_paragraph(doc, "BY", size=12)
    add_center_paragraph(doc, STUDENT_NAME, bold=True, size=12)
    add_center_paragraph(doc, ROLL_NO, bold=True, size=12, space_after=12)
    add_center_paragraph(doc, "Under the Guidance of", size=12)
    add_center_paragraph(doc, GUIDES, bold=True, size=12, space_after=12)
    add_center_paragraph(doc, "Department of Computer Science & Engineering", size=12)
    add_center_paragraph(doc, "Faculty of Engineering", size=12)
    add_center_paragraph(doc, "MEDICAPS UNIVERSITY, INDORE - 453331", bold=True, size=12)
    add_center_paragraph(doc, MONTH_YEAR, bold=True, size=12)

    add_page_break(doc)

    # Front matter
    add_heading_line(doc, "Approval Sheet")
    add_body_line(doc, f"The project work titled '{TITLE}' is approved as a creditable study in partial fulfillment of the degree requirements.")
    add_body_line(doc, "Internal Examiner: ____________________")
    add_body_line(doc, "External Examiner: ____________________")

    add_page_break(doc)
    add_heading_line(doc, "Declaration")
    add_body_line(doc, f"I hereby declare that the project titled '{TITLE}' submitted in partial fulfillment for the award of the degree of Bachelor of Technology in {DEPARTMENT} has been carried out by me under the guidance of {GUIDES}, Faculty of Engineering, Medi-Caps University, Indore.")
    add_body_line(doc, "I further declare that this work is original and has not been submitted earlier for any other degree or diploma.")
    add_body_line(doc, f"Signature: {STUDENT_NAME}")
    add_body_line(doc, f"Date: {date.today().isoformat()}")

    add_page_break(doc)
    add_heading_line(doc, "Certificate")
    add_body_line(doc, f"This is to certify that the project report titled '{TITLE}' submitted by {STUDENT_NAME} ({ROLL_NO}) is a bonafide record of work carried out under our supervision.")
    add_body_line(doc, f"Guide(s): {GUIDES}")
    add_body_line(doc, "Head of Department: ____________________")

    add_page_break(doc)
    add_heading_line(doc, "Acknowledgement")
    add_body_line(doc, "I express my sincere gratitude to my guides for their consistent support, technical direction, and encouragement throughout this project.")
    add_body_line(doc, "I also thank the Department of Computer Science and Engineering, Faculty of Engineering, Medi-Caps University, for providing the academic environment and resources necessary to complete this work.")
    add_body_line(doc, "I am thankful to my peers and family members for their motivation and support during design, implementation, and testing stages.")

    add_page_break(doc)
    add_heading_line(doc, "Table of Contents")
    add_body_line(doc, "Abstract")
    add_body_line(doc, "Chapter 1: Introduction")
    add_body_line(doc, "Chapter 2: Requirements Specification")
    add_body_line(doc, "Chapter 3: System Design")
    add_body_line(doc, "Chapter 4: Implementation, Testing, and Maintenance")
    add_body_line(doc, "Chapter 5: Results and Discussion")
    add_body_line(doc, "Chapter 6: Summary and Conclusions")
    add_body_line(doc, "Chapter 7: Future Scope")
    add_body_line(doc, "Bibliography")

    # Main chapters
    for heading, lines in SECTIONS:
        add_page_break(doc)
        add_heading_line(doc, heading)
        for line in lines:
            if line.startswith("Chapter ") or line[:3].isdigit() or (len(line) > 3 and line[0].isdigit() and line[1] == "."):
                add_heading_line(doc, line)
            else:
                add_body_line(doc, line)

    doc.save(OUTPUT_FILE)
    print(OUTPUT_FILE)


if __name__ == "__main__":
    build_report()
