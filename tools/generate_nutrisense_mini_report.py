from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_PATH = Path(r"C:\Users\DELL\Downloads\BTech_CSE_MINI Project- Report Format.docx")
OUTPUT_PATH = ROOT / "NutriSense_Mini_Project_Report_Final.docx"

TITLE = "NUTRISENSE: PERSONALIZED DIET AND FITNESS PLANNER FOR INDIAN USERS"
STUDENT_NAME = "Prabhmeet Singh"
ROLL_NO = "EN23CS301742"
GUIDE_LINE_1 = "Mr. Vishal Sharma"
GUIDE_LINE_2 = "Ms. Anusha Jain"
DEPARTMENT = "Department of Computer Science & Engineering"
INSTITUTE = "MEDICAPS UNIVERSITY, INDORE - 453331"
MONTH_YEAR = "APRIL 2026"


def clear_document(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_page_format(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Pt(30)
    section.bottom_margin = Pt(22)
    section.left_margin = Pt(30 * 2.83465)
    section.right_margin = Pt(20 * 2.83465)


def set_default_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)


def add_center(doc: Document, text: str, bold: bool = False, size: int = 12, after: int = 8) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.5


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(16 if level == 1 else 14)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.5


def add_subheading(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Heading 2")
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5


def add_para(doc: Document, text: str, indent: bool = True) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.first_line_indent = Pt(18)


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def add_field(paragraph, instruction: str, display_text: str) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    txt = OxmlElement("w:t")
    txt.text = display_text

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(txt)
    run._r.append(fld_end)


def add_seq_caption(paragraph, label: str, title: str) -> None:
    paragraph.style = "Caption"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(f"{label} ")
    add_field(paragraph, f'SEQ {label} \\* ARABIC', "1")
    paragraph.add_run(f": {title}")


def add_toc_field(doc: Document, title: str, instruction: str, placeholder: str) -> None:
    add_heading(doc, title, level=1)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    add_field(p, instruction, placeholder)
    add_para(doc, "Note: Press F9 in Microsoft Word to update this generated list.", indent=False)


def add_table(doc: Document, heading: str, columns: list[str], rows: list[list[str]], table_no: str) -> None:
    add_subheading(doc, heading)
    cap = doc.add_paragraph()
    add_seq_caption(cap, "Table", heading)
    cap.paragraph_format.space_after = Pt(4)

    table = doc.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0].cells
    for i, col in enumerate(columns):
        hdr[i].text = col

    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val

    doc.add_paragraph()


def add_figure(doc: Document, image_rel: str, caption: str) -> None:
    image_path = ROOT / image_rel
    if not image_path.exists():
        add_para(doc, f"[Figure missing: {caption}]", indent=False)
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(6.1))

    cap = doc.add_paragraph()
    add_seq_caption(cap, "Figure", caption)
    cap.paragraph_format.space_after = Pt(8)


def add_front_cover(doc: Document) -> None:
    add_center(doc, TITLE, bold=True, size=18, after=26)
    add_center(doc, "A Minor Project Report", bold=True, size=14)
    add_center(doc, "Submitted in partial fulfillment of requirement of the", size=12)
    add_center(doc, "Degree of", size=12)
    add_center(doc, "BACHELOR OF TECHNOLOGY in COMPUTER SCIENCE & ENGINEERING", bold=True, size=13, after=16)
    add_center(doc, "BY", bold=True, size=12)
    add_center(doc, STUDENT_NAME, bold=True, size=13)
    add_center(doc, ROLL_NO, bold=True, size=12, after=16)
    add_center(doc, "Under the Guidance of", size=12)
    add_center(doc, GUIDE_LINE_1, bold=True, size=12)
    add_center(doc, GUIDE_LINE_2, bold=True, size=12, after=16)
    add_center(doc, DEPARTMENT, size=12)
    add_center(doc, "Faculty of Engineering", size=12)
    add_center(doc, INSTITUTE, bold=True, size=12)
    add_center(doc, MONTH_YEAR, bold=True, size=12)


def add_title_page(doc: Document) -> None:
    add_center(doc, TITLE, bold=True, size=16, after=20)
    add_center(doc, "A Minor Project Report", bold=True, size=13)
    add_center(doc, "Submitted in partial fulfillment of requirement of the", size=12)
    add_center(doc, "Degree of", size=12)
    add_center(doc, "BACHELOR OF TECHNOLOGY in COMPUTER SCIENCE & ENGINEERING", bold=True, size=12, after=14)
    add_center(doc, "BY", bold=True, size=12)
    add_center(doc, STUDENT_NAME, bold=True, size=12)
    add_center(doc, ROLL_NO, bold=True, size=12, after=14)
    add_center(doc, "Under the Guidance of", size=12)
    add_center(doc, GUIDE_LINE_1 + " and " + GUIDE_LINE_2, bold=True, size=12, after=14)
    add_center(doc, DEPARTMENT, size=12)
    add_center(doc, "Faculty of Engineering", size=12)
    add_center(doc, INSTITUTE, bold=True, size=12)
    add_center(doc, MONTH_YEAR, bold=True, size=12)


def add_approval(doc: Document) -> None:
    add_heading(doc, "Report Approval", level=1)
    add_para(
        doc,
        f"The project work titled '{TITLE}' is hereby approved as a creditable study of an engineering application subject carried out and presented in a satisfactory manner for partial fulfillment of the B.Tech degree requirements.",
    )
    add_para(
        doc,
        "It is to be understood that by this approval, the undersigned do not endorse any statement or conclusion in the report beyond its academic acceptability.",
    )
    doc.add_paragraph("Internal Examiner: __________________________")
    doc.add_paragraph("External Examiner: __________________________")


def add_declaration(doc: Document) -> None:
    add_heading(doc, "Declaration", level=1)
    add_para(
        doc,
        f"I hereby declare that the project entitled '{TITLE}' submitted in partial fulfillment for the award of Bachelor of Technology in Computer Science and Engineering is an authentic work carried out by me under the supervision of {GUIDE_LINE_1} and {GUIDE_LINE_2}, Faculty of Engineering, Medi-Caps University, Indore.",
    )
    add_para(
        doc,
        "I further declare that this project work, in full or in part, has neither been copied from any other source nor submitted to any other institute or university for any degree or diploma.",
    )
    doc.add_paragraph(f"Signature: {STUDENT_NAME}")
    doc.add_paragraph(f"Date: {date.today().strftime('%d-%m-%Y')}")


def add_certificate(doc: Document) -> None:
    add_heading(doc, "Certificate", level=1)
    add_para(
        doc,
        f"This is to certify that the project report titled '{TITLE}' submitted by {STUDENT_NAME} ({ROLL_NO}) in partial fulfillment of the requirements for the degree of Bachelor of Technology in Computer Science and Engineering is a bonafide record of work carried out under our guidance.",
    )
    doc.add_paragraph("_______________________________")
    doc.add_paragraph(GUIDE_LINE_1)
    doc.add_paragraph("Project Guide")
    doc.add_paragraph("_______________________________")
    doc.add_paragraph(GUIDE_LINE_2)
    doc.add_paragraph("Project Co-Guide")
    doc.add_paragraph("_______________________________")
    doc.add_paragraph("Dr. Ratnesh Litoriya")
    doc.add_paragraph("Head, Department of Computer Science & Engineering")


def add_acknowledgement(doc: Document) -> None:
    add_heading(doc, "Acknowledgement", level=1)
    add_para(
        doc,
        "I express my sincere gratitude to my project guides, Mr. Vishal Sharma and Ms. Anusha Jain, for their valuable guidance, consistent encouragement, and constructive feedback during all phases of this project.",
    )
    add_para(
        doc,
        "I am grateful to the Department of Computer Science and Engineering, Faculty of Engineering, Medi-Caps University, for providing the academic environment, infrastructure, and support needed for successful completion of this work.",
    )
    add_para(
        doc,
        "I also acknowledge the support of my peers and family members whose motivation and cooperation helped me complete implementation, testing, and documentation in a timely manner.",
    )
    doc.add_paragraph(STUDENT_NAME)
    doc.add_paragraph("B.Tech. III Year")


def add_abstract(doc: Document) -> None:
    add_heading(doc, "Abstract", level=1)
    add_para(
        doc,
        "NutriSense is a full-stack web application that generates personalized diet and fitness plans for Indian users using body profile, activity level, diet preference, and target goals. The project addresses key limitations in generic health apps by integrating Indian meal patterns, affordability-aware recommendations, and beginner-friendly weekly workout schedules in one unified interface.",
    )
    add_para(
        doc,
        "The backend is built with FastAPI and modular services. Diet generation uses rule-based computation including BMI estimation, goal mapping, TDEE-based calorie planning, and protein distribution across four meals: breakfast, lunch, evening meal, and dinner. Fitness generation provides structured seven-day schedules with focus-wise exercises, duration guidance, and recovery recommendations. Frontend implementation in React with TypeScript provides an interactive user flow from profile capture to plan visualization.",
    )
    add_para(
        doc,
        "System validation across multiple user scenarios shows deterministic and complete output generation for different goals and diet types. The platform provides macro summaries, budget hints, itemized meal details, and explainable recommendations suitable for students and working professionals. NutriSense demonstrates a practical and extensible architecture for personalized digital wellness planning in the Indian context.",
    )
    add_para(
        doc,
        "Keywords: personalized nutrition, Indian diet planning, BMI, TDEE, FastAPI, React, workout scheduling, budget-aware recommendations",
        indent=False,
    )


def add_abbreviations(doc: Document) -> None:
    add_heading(doc, "Abbreviations", level=1)
    entries = [
        ("API", "Application Programming Interface"),
        ("BMI", "Body Mass Index"),
        ("BMR", "Basal Metabolic Rate"),
        ("TDEE", "Total Daily Energy Expenditure"),
        ("UI", "User Interface"),
        ("UX", "User Experience"),
        ("JSON", "JavaScript Object Notation"),
        ("REST", "Representational State Transfer"),
        ("IDE", "Integrated Development Environment"),
        ("MUI", "Material UI"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Abbreviation"
    table.rows[0].cells[1].text = "Expanded Form"
    for abbr, exp in entries:
        cells = table.add_row().cells
        cells[0].text = abbr
        cells[1].text = exp


def add_notations(doc: Document) -> None:
    add_heading(doc, "Notations & Symbols", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Notation"
    table.rows[0].cells[1].text = "Meaning"

    notations = [
        ("BMI", "Body mass index, computed as weight/height^2"),
        ("W", "Body weight in kilograms"),
        ("H", "Height in meters"),
        ("P_d", "Daily protein requirement (grams/day)"),
        ("C_t", "Daily calorie target (kcal/day)"),
        ("BMR", "Basal metabolic rate (kcal/day)"),
        ("TDEE", "Total daily energy expenditure (kcal/day)"),
        ("M_b, M_l, M_e, M_d", "Breakfast, lunch, evening, and dinner macro allocations"),
    ]

    for symbol, meaning in notations:
        row = table.add_row().cells
        row[0].text = symbol
        row[1].text = meaning

    add_para(doc, "Core equations used in recommendation engine:", indent=False)
    add_para(doc, "BMI = W / H^2", indent=False)
    add_para(doc, "BMR = 10W + 6.25H_cm - 5A + 5", indent=False)
    add_para(doc, "TDEE = BMR x activity_multiplier", indent=False)
    add_para(doc, "P_d = W x k, where k belongs to {1.2, 1.5, 1.8}", indent=False)


def add_chapter_1(doc: Document) -> None:
    add_heading(doc, "Chapter 1", level=1)
    add_heading(doc, "INTRODUCTION", level=1)

    add_subheading(doc, "1.1 Introduction")
    add_para(
        doc,
        "Personalized health planning is difficult for students and early-career professionals due to inconsistent schedules, limited budgets, and lack of localized dietary guidance. Most popular applications provide global recommendations that do not align with Indian food habits. NutriSense addresses this practical gap by generating implementable diet and fitness plans from basic user inputs.",
    )

    add_subheading(doc, "1.2 Literature Review")
    add_para(
        doc,
        "Recent studies indicate that personalized nutrition improves adherence and outcomes over generic plans [1], [5]. BMI-based classification is widely used for initial risk stratification [3], while TDEE-based energy estimation supports practical calorie planning [2]. Existing systems often optimize for broader markets and miss Indian meal localization and affordability constraints. NutriSense focuses on these two limitations through rule-based recommendation and budget-aware output.",
    )

    add_subheading(doc, "1.3 Objectives")
    add_para(doc, "The major objectives of NutriSense are listed below:")
    objectives = [
        "To compute BMI, calorie, and protein targets from user profile data.",
        "To generate complete four-meal daily plans with macros and cost ranges.",
        "To support diet types including vegetarian, non-vegetarian, vegan, and eggetarian.",
        "To generate weekly fitness routines aligned with user goals.",
        "To present a clean, explainable, and user-friendly web interface.",
    ]
    for item in objectives:
        add_para(doc, f"- {item}", indent=False)

    add_subheading(doc, "1.4 Significance")
    add_para(
        doc,
        "The project is significant because it delivers a practical planning system for Indian users rather than only a theoretical recommendation engine. By combining meal planning and fitness guidance with budget sensitivity, NutriSense improves real-world adoptability.",
    )

    add_subheading(doc, "1.5 Research Design")
    add_para(
        doc,
        "An iterative design-and-build method was used. Core formulas and rule sets were defined first, followed by backend module implementation, API integration, frontend rendering, and scenario-based testing. Each iteration validated output completeness, consistency, and usability.",
    )

    add_subheading(doc, "1.6 Source of Data")
    add_para(
        doc,
        "Food options and macro estimates are maintained in Python model dictionaries curated for common Indian meals. User profile values are collected from frontend forms. Fitness plans are generated through structured exercise templates based on goal and experience level.",
    )

    add_subheading(doc, "1.7 Chapter Scheme")
    add_para(
        doc,
        "Chapter 2 presents requirements specification. Chapter 3 explains system design with diagrams and architecture. Chapter 4 details implementation and testing. Chapter 5 discusses results and observations. Chapter 6 summarizes conclusions and Chapter 7 outlines future scope.",
    )


def add_chapter_2(doc: Document) -> None:
    add_heading(doc, "Chapter 2", level=1)
    add_heading(doc, "REQUIREMENTS SPECIFICATION", level=1)

    add_subheading(doc, "2.1 User Characteristics")
    add_para(
        doc,
        "Target users include students, hostel residents, and working professionals looking for practical meal and workout plans without expensive subscriptions or complex onboarding. Users may have varying fitness backgrounds and dietary preferences.",
    )

    add_subheading(doc, "2.2 Functional Requirements")
    add_table(
        doc,
        "Functional Requirements",
        ["ID", "Requirement", "Description"],
        [
            ["FR-1", "User Input Capture", "Collect name, age, height, weight, goal, fitness level, and diet type."],
            ["FR-2", "BMI and Goal Mapping", "Calculate BMI and infer default recommendation context."],
            ["FR-3", "Meal Plan Generation", "Generate breakfast, lunch, evening, and dinner recommendations."],
            ["FR-4", "Fitness Schedule", "Generate a seven-day workout schedule with intensity mapping."],
            ["FR-5", "Results Dashboard", "Display macro details, tips, hydration, and cost ranges in UI."],
        ],
        "2.1",
    )

    add_subheading(doc, "2.3 Dependencies")
    add_para(
        doc,
        "Backend dependencies include FastAPI, Uvicorn, Pydantic, and pymongo. Frontend dependencies include React, TypeScript, Vite, MUI, axios, and react-router-dom.",
    )

    add_subheading(doc, "2.4 Performance Requirements")
    add_para(
        doc,
        "The system should return complete recommendations in interactive time for normal user inputs. Equivalent inputs should produce deterministic outputs under the same rule configuration.",
    )

    add_subheading(doc, "2.5 Hardware Requirements")
    add_para(
        doc,
        "Minimum hardware target is a modern laptop/desktop with 8 GB RAM and internet-enabled browser for frontend access. Development setup requires Python 3.11 and Node.js runtime.",
    )

    add_subheading(doc, "2.6 Constraints & Assumptions")
    add_para(
        doc,
        "Recommendations are intended for wellness guidance and are not a medical diagnosis. Nutritional values and costs are approximate. User-provided measurements are assumed to be correct.",
    )


def add_chapter_3(doc: Document) -> None:
    add_heading(doc, "Chapter 3", level=1)
    add_heading(doc, "DESIGN", level=1)

    add_subheading(doc, "3.1 Algorithm")
    add_para(
        doc,
        "NutriSense combines BMI calculation, goal derivation, TDEE computation, and protein multipliers. Daily protein is distributed using the ratio 25:30:20:25 across breakfast, lunch, evening, and dinner. Calorie targets are adjusted using goal context, followed by diet-type template mapping.",
    )

    add_subheading(doc, "3.2 Function Oriented Design")
    add_para(
        doc,
        "The backend is modularized into API routing, diet service, fitness service, and food model modules. The API receives user payload, orchestrates service calls, and returns a unified JSON response to the frontend.",
    )

    add_subheading(doc, "3.3 System Design")
    add_subheading(doc, "3.3.1 Data Flow Diagrams")
    add_figure(doc, "docs/sds-diagrams/dfd.png", "Figure 3.1: Data Flow Diagram")

    add_subheading(doc, "3.3.2 Activity Diagram")
    add_figure(doc, "docs/sds-diagrams/activity_diagram.png", "Figure 3.2: Activity Diagram")

    add_subheading(doc, "3.3.3 Flow Chart")
    add_figure(doc, "docs/sds-diagrams/flowchart.png", "Figure 3.3: Flowchart of Recommendation Workflow")

    add_subheading(doc, "3.3.4 Class Diagram")
    add_figure(doc, "docs/sds-diagrams/class_diagram.png", "Figure 3.4: Class Diagram")

    add_subheading(doc, "3.3.5 ER Diagram")
    add_figure(doc, "docs/sds-diagrams/er_diagram.png", "Figure 3.5: ER Diagram")

    add_subheading(doc, "3.3.6 Sequence Diagram")
    add_figure(doc, "docs/sds-diagrams/sequence_diagram.png", "Figure 3.6: Sequence Diagram")

    add_subheading(doc, "3.4 Database Design")
    add_subheading(doc, "3.4.1 Logical Database Design")
    add_para(
        doc,
        "Current implementation uses MongoDB collections for users, plan_history, and plan_feedback. Input and generated recommendation payloads are stored with timestamps for retrieval and feedback analysis.",
    )
    add_subheading(doc, "3.4.2 Physical Database Design")
    add_para(
        doc,
        "Collections are structured for lightweight JSON documents. Primary retrieval keys include email, plan_id, and created_at. The design is optimized for iterative project development and easy extension.",
    )


def add_chapter_4(doc: Document) -> None:
    add_heading(doc, "Chapter 4", level=1)
    add_heading(doc, "IMPLEMENTATION, TESTING, AND MAINTENANCE", level=1)

    add_subheading(doc, "4.1 Introduction to Languages, IDEs, Tools and Technologies")
    add_table(
        doc,
        "Technology Stack Used in NutriSense",
        ["Layer", "Technology", "Purpose"],
        [
            ["Frontend", "React + TypeScript + Vite", "User interface and routing"],
            ["Backend", "FastAPI (Python)", "REST API and recommendation engine"],
            ["Database", "MongoDB", "User and plan history storage"],
            ["HTTP Client", "axios", "Frontend-backend API communication"],
            ["UI Library", "MUI", "Responsive components and theming"],
            ["Development", "VS Code", "Coding, debugging, and integration"],
        ],
        "4.1",
    )

    add_subheading(doc, "4.2 Testing Techniques and Test Plans")
    add_para(
        doc,
        "Testing included endpoint validation, schema checks, and scenario-based functional testing for all supported goals and diet types. UI testing verified that each generated meal card and weekly schedule renders correctly.",
    )
    add_table(
        doc,
        "Representative Functional Test Cases",
        ["Test ID", "Input Scenario", "Expected Outcome", "Status"],
        [
            ["TC-01", "Beginner + veg + maintenance", "4 meals + weekly plan generated", "Pass"],
            ["TC-02", "Intermediate + non_veg + muscle_gain", "Higher protein plan with workout split", "Pass"],
            ["TC-03", "Advanced + vegan + weight_loss", "Calorie-adjusted plan and recovery tips", "Pass"],
            ["TC-04", "Invalid payload values", "Validation error response", "Pass"],
        ],
        "4.2",
    )

    add_subheading(doc, "4.3 Installation Instructions")
    add_para(doc, "Backend setup:", indent=False)
    add_para(doc, "1. Navigate to DietAndFitnessPlanner-BE.", indent=False)
    add_para(doc, "2. Create/activate virtual environment and install dependencies from requirements.txt.", indent=False)
    add_para(doc, "3. Run the server using: uvicorn app.main:app --reload.", indent=False)
    add_para(doc, "Frontend setup:", indent=False)
    add_para(doc, "1. Navigate to DietAndFitnessPlanner-FE.", indent=False)
    add_para(doc, "2. Install dependencies using npm install.", indent=False)
    add_para(doc, "3. Run using npm run dev and open the provided localhost URL.", indent=False)

    add_subheading(doc, "4.4 End User Instructions")
    add_para(doc, "1. Open NutriSense web application.", indent=False)
    add_para(doc, "2. Log in and enter profile details in the form.", indent=False)
    add_para(doc, "3. Select fitness level, goal, diet type, and budget preference.", indent=False)
    add_para(doc, "4. Click Generate Plan to view personalized output.", indent=False)
    add_para(doc, "5. Review macro totals, meal breakdown, and weekly workout details.", indent=False)


def add_chapter_5(doc: Document) -> None:
    add_heading(doc, "Chapter 5", level=1)
    add_heading(doc, "RESULTS AND DISCUSSIONS", level=1)

    add_subheading(doc, "5.1 User Interface Representation")
    add_para(
        doc,
        "The user journey covers landing page, login page, details form, and result dashboard. The final dashboard includes overview cards, diet tabs, and workout schedule panels.",
    )

    add_subheading(doc, "5.2 Brief Description of Various Modules")
    add_para(
        doc,
        "Diet module computes BMI, protein targets, macro distribution, and meal recommendations. Fitness module maps user goal to workout type, applies intensity levels, and generates seven-day plans with recovery suggestions. API router combines outputs and handles persistence for plan history and feedback.",
    )

    add_subheading(doc, "5.3 Snapshots of System with Brief Details")
    add_figure(doc, "docs/sds-diagrams/system_architecture.png", "Figure 5.1: System Architecture Overview")
    add_figure(doc, "docs/sds-diagrams/component_diagram.png", "Figure 5.2: Component Diagram")
    add_figure(doc, "docs/sds-diagrams/deployment_diagram.png", "Figure 5.3: Deployment Diagram")

    add_subheading(doc, "5.4 Back End Representation")
    add_table(
        doc,
        "Sample API Output Summary",
        ["Field", "Sample Value", "Description"],
        [
            ["status", "success", "Request processing status"],
            ["diet_plan.bmi", "22.8", "Computed BMI"],
            ["diet_plan.meal_plan", "4 meals", "Breakfast, lunch, evening, dinner"],
            ["fitness_plan.detailed_plan", "7 days", "Weekly workout structure"],
            ["estimated_daily_cost_range", "{min,max}", "Daily spend guidance"],
        ],
        "5.1",
    )

    add_subheading(doc, "5.5 Snapshots of Database Tables with Brief Description")
    add_table(
        doc,
        "MongoDB Collections in Current Build",
        ["Collection", "Key Fields", "Purpose"],
        [
            ["users", "email, name, age, updated_at", "User profile persistence"],
            ["plan_history", "email, goal, diet_plan, fitness_plan, created_at", "Generated plan audit trail"],
            ["plan_feedback", "plan_id, adherence_percent, feedback_text", "User adherence and feedback capture"],
        ],
        "5.2",
    )

    add_para(
        doc,
        "Discussion: The implementation demonstrates that a rule-based localized planner can provide complete and explainable outputs with low computational complexity. The architecture supports future extension toward data-driven personalization while retaining transparency.",
    )


def add_chapter_6(doc: Document) -> None:
    add_heading(doc, "Chapter 6", level=1)
    add_heading(doc, "SUMMARY AND CONCLUSIONS", level=1)
    add_para(
        doc,
        "NutriSense successfully integrates personalized diet generation and weekly fitness planning in a single web-based platform. The system meets the intended objectives of localized recommendation, budget-aware guidance, and user-friendly output visualization.",
    )
    add_para(
        doc,
        "From an engineering perspective, the project demonstrates a modular FastAPI + React architecture with clear service boundaries, maintainable code flow, and scalable extension points. Rule-based logic provides deterministic and explainable behavior, which is beneficial for educational and prototype healthcare systems.",
    )
    add_para(
        doc,
        "Overall, NutriSense is a practical, deployable foundation for personalized wellness support for Indian users.",
    )


def add_chapter_7(doc: Document) -> None:
    add_heading(doc, "Chapter 7", level=1)
    add_heading(doc, "FUTURE SCOPE", level=1)
    points = [
        "Integration of real-time wearable data for adaptive recommendations.",
        "Inclusion of multilingual support for wider regional adoption.",
        "Expansion of food datasets with dynamic substitutions and allergy filters.",
        "Advanced analytics dashboard for long-term progress tracking.",
        "Hybrid rule + machine learning recommender with safety constraints.",
    ]
    for item in points:
        add_para(doc, f"- {item}", indent=False)


def add_appendix(doc: Document) -> None:
    add_heading(doc, "Appendix", level=1)
    add_subheading(doc, "Appendix A: Formula Mapping")
    add_para(doc, "BMI thresholds: <18.5 (muscle gain), 18.5-25 (maintenance), >25 (weight loss).", indent=False)
    add_para(doc, "Protein multipliers: beginner=1.2 g/kg, intermediate=1.5 g/kg, advanced=1.8 g/kg.", indent=False)
    add_para(doc, "Meal protein split: breakfast 25%, lunch 30%, evening 20%, dinner 25%.", indent=False)

    add_subheading(doc, "Appendix B: Budget Slab Mapping")
    add_para(doc, "Affordable: Rs. 100-250/day", indent=False)
    add_para(doc, "Value/Balanced: Rs. 250-350/day", indent=False)
    add_para(doc, "Premium: Rs. 350-500/day", indent=False)


def add_bibliography(doc: Document) -> None:
    add_heading(doc, "Bibliography", level=1)
    refs = [
        "[1] World Health Organization, Obesity and Overweight, WHO Fact Sheets, 2024.",
        "[2] M. D. Mifflin, S. T. St Jeor, et al., A new predictive equation for resting energy expenditure in healthy individuals, The American Journal of Clinical Nutrition, vol. 51, no. 2, pp. 241-247, 1990.",
        "[3] National Institutes of Health, Clinical Guidelines on the Identification, Evaluation, and Treatment of Overweight and Obesity in Adults, NIH Publication, 1998.",
        "[4] Harvard T.H. Chan School of Public Health, The Nutrition Source: Healthy Eating Plate, 2023.",
        "[5] C. Celis-Morales, K. Livingstone, et al., Personalized nutrition and health outcomes: A review of approaches and evidence, Proceedings of the Nutrition Society, 2021.",
        "[6] Ministry of Health and Family Welfare, Government of India, Diet and Wellness Guidelines for Adults, 2022.",
    ]
    for ref in refs:
        add_para(doc, ref, indent=False)


def add_publications(doc: Document) -> None:
    add_heading(doc, "List of Publications", level=1)
    add_para(doc, "No publications from this project at the time of report submission.", indent=False)

    add_heading(doc, "Reprints of Publications", level=1)
    add_para(doc, "Not applicable.", indent=False)


def build_report() -> Path:
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Template not found: {TEMPLATE_PATH}")

    doc = Document(str(TEMPLATE_PATH))
    clear_document(doc)
    set_page_format(doc)
    set_default_style(doc)

    add_front_cover(doc)
    add_page_break(doc)
    add_title_page(doc)

    add_page_break(doc)
    add_approval(doc)
    add_page_break(doc)
    add_declaration(doc)
    add_page_break(doc)
    add_certificate(doc)
    add_page_break(doc)
    add_acknowledgement(doc)
    add_page_break(doc)
    add_abstract(doc)

    add_page_break(doc)
    add_toc_field(doc, "Table of Contents", 'TOC \\o "1-3" \\h \\z \\u', "Table of Contents (update field)")
    add_page_break(doc)
    add_toc_field(doc, "List of Figures", 'TOC \\h \\z \\c "Figure"', "List of Figures (update field)")
    add_page_break(doc)
    add_toc_field(doc, "List of Tables", 'TOC \\h \\z \\c "Table"', "List of Tables (update field)")
    add_page_break(doc)
    add_abbreviations(doc)
    add_page_break(doc)
    add_notations(doc)

    add_page_break(doc)
    add_chapter_1(doc)
    add_page_break(doc)
    add_chapter_2(doc)
    add_page_break(doc)
    add_chapter_3(doc)
    add_page_break(doc)
    add_chapter_4(doc)
    add_page_break(doc)
    add_chapter_5(doc)
    add_page_break(doc)
    add_chapter_6(doc)
    add_page_break(doc)
    add_chapter_7(doc)
    add_page_break(doc)
    add_appendix(doc)
    add_page_break(doc)
    add_bibliography(doc)
    add_page_break(doc)
    add_publications(doc)

    doc.save(str(OUTPUT_PATH))
    return OUTPUT_PATH


if __name__ == "__main__":
    out = build_report()
    print(out)