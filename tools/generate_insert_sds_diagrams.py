from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List, Tuple

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Inches


W, H = 1800, 1100
BG = "#F8FAFC"
PRIMARY = "#1F3A5F"
SECONDARY = "#2E6EA6"
ACCENT = "#2A9D8F"
TEXT = "#0F172A"
BOX_FILL = "#FFFFFF"
BOX_BORDER = "#2E6EA6"


@dataclass
class Box:
    text: str
    x: int
    y: int
    w: int = 300
    h: int = 120
    fill: str = BOX_FILL


def get_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "C:/Windows/Fonts/segoeui.ttf",
        "C:/Windows/Fonts/arial.ttf",
    ]
    if bold:
        candidates = [
            "C:/Windows/Fonts/segoeuib.ttf",
            "C:/Windows/Fonts/arialbd.ttf",
            *candidates,
        ]
    for path in candidates:
        p = Path(path)
        if p.exists():
            return ImageFont.truetype(str(p), size=size)
    return ImageFont.load_default()


FONT_TITLE = get_font(46, bold=True)
FONT_LABEL = get_font(28, bold=True)
FONT_TEXT = get_font(24)
FONT_SMALL = get_font(20)


def draw_header(draw: ImageDraw.ImageDraw, title: str, subtitle: str) -> None:
    draw.rectangle([0, 0, W, 150], fill=PRIMARY)
    draw.text((60, 35), title, font=FONT_TITLE, fill="white")
    draw.text((62, 95), subtitle, font=FONT_SMALL, fill="#D8E7F6")


def draw_box(draw: ImageDraw.ImageDraw, box: Box) -> None:
    x1, y1 = box.x, box.y
    x2, y2 = box.x + box.w, box.y + box.h
    draw.rounded_rectangle([x1, y1, x2, y2], radius=22, fill=box.fill, outline=BOX_BORDER, width=4)
    lines = wrap_text(draw, box.text, FONT_TEXT, box.w - 30)
    total_h = len(lines) * 30
    y = box.y + (box.h - total_h) // 2
    for line in lines:
        tw = draw.textlength(line, font=FONT_TEXT)
        draw.text((box.x + (box.w - tw) / 2, y), line, font=FONT_TEXT, fill=TEXT)
        y += 30


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_w: int) -> List[str]:
    words = text.split()
    lines: List[str] = []
    current: List[str] = []
    for word in words:
        trial = " ".join(current + [word])
        if draw.textlength(trial, font=font) <= max_w:
            current.append(word)
        else:
            if current:
                lines.append(" ".join(current))
            current = [word]
    if current:
        lines.append(" ".join(current))
    return lines or [text]


def arrow(draw: ImageDraw.ImageDraw, p1: Tuple[int, int], p2: Tuple[int, int], color: str = SECONDARY, width: int = 5) -> None:
    draw.line([p1, p2], fill=color, width=width)
    dx = p2[0] - p1[0]
    dy = p2[1] - p1[1]
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    hx, hy = p2
    left = (hx - 20 * ux + 10 * uy, hy - 20 * uy - 10 * ux)
    right = (hx - 20 * ux - 10 * uy, hy - 20 * uy + 10 * ux)
    draw.polygon([p2, left, right], fill=color)


def make_canvas(title: str, subtitle: str) -> Tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", (W, H), BG)
    draw = ImageDraw.Draw(img)
    draw_header(draw, title, subtitle)
    return img, draw


def save(img: Image.Image, out_path: Path) -> None:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    img.save(out_path, format="PNG")


def build_use_case(path: Path) -> None:
    img, draw = make_canvas("Use Case Diagram", "NutriSense - Personalized Diet and Fitness Planner")
    draw.ellipse([650, 200, 1500, 980], outline=SECONDARY, width=5)
    actor = Box("User", 120, 450, 180, 120, fill="#E8F3FF")
    draw_box(draw, actor)
    cases = [
        Box("Register / Login", 760, 260),
        Box("Enter Profile Details", 980, 400),
        Box("Generate Diet Plan", 760, 540),
        Box("Generate Fitness Plan", 980, 680),
        Box("View Recommendations", 760, 820),
    ]
    for c in cases:
        draw_box(draw, c)
        arrow(draw, (300, 510), (c.x, c.y + c.h // 2))
    save(img, path)


def build_dfd(path: Path) -> None:
    img, draw = make_canvas("Data Flow Diagram (Level 1)", "NutriSense Data Processing Flow")
    ext = Box("User", 90, 420, 220, 120, fill="#E8F3FF")
    p1 = Box("Input Validation", 430, 280)
    p2 = Box("Diet Service", 860, 200)
    p3 = Box("Fitness Service", 860, 430)
    ds1 = Box("Food Dataset", 1300, 170, 320, 130, fill="#ECFDF5")
    ds2 = Box("Workout Rules", 1300, 430, 320, 130, fill="#ECFDF5")
    p4 = Box("Plan Composer", 860, 700)
    out = Box("Result UI", 1300, 760, 320, 130, fill="#E8F3FF")
    for b in [ext, p1, p2, p3, ds1, ds2, p4, out]:
        draw_box(draw, b)
    arrow(draw, (310, 480), (430, 340))
    arrow(draw, (730, 340), (860, 260))
    arrow(draw, (730, 340), (860, 490))
    arrow(draw, (1160, 260), (1300, 235))
    arrow(draw, (1160, 490), (1300, 495))
    arrow(draw, (1010, 330), (1010, 700))
    arrow(draw, (1010, 560), (1010, 700))
    arrow(draw, (1160, 760), (1300, 825))
    save(img, path)


def build_er(path: Path) -> None:
    img, draw = make_canvas("ER Diagram", "Logical Data Entities for NutriSense")
    entities = [
        Box("UserProfile\nuser_id (PK)\nage\nheight\nweight", 120, 260, 360, 220),
        Box("NutritionTarget\ntarget_id (PK)\ncalories\nprotein\nuser_id (FK)", 700, 180, 380, 230),
        Box("MealPlan\nmeal_plan_id (PK)\nmeal_type\nmacro_split\nuser_id (FK)", 1320, 180, 380, 230),
        Box("FitnessPlan\nfitness_plan_id (PK)\nweek_schedule\nintensity\nuser_id (FK)", 700, 560, 380, 230),
        Box("RecommendationResult\nresult_id (PK)\ncreated_at\nuser_id (FK)", 1320, 560, 380, 210),
    ]
    for e in entities:
        draw_box(draw, e)
    arrow(draw, (480, 360), (700, 295), color=ACCENT)
    arrow(draw, (1080, 295), (1320, 295), color=ACCENT)
    arrow(draw, (480, 390), (700, 675), color=ACCENT)
    arrow(draw, (1080, 675), (1320, 665), color=ACCENT)
    save(img, path)


def build_flow(path: Path) -> None:
    img, draw = make_canvas("Flowchart", "End-to-End Recommendation Workflow")
    steps = [
        Box("Start", 760, 180, 280, 90, fill="#ECFDF5"),
        Box("Login / Authenticate", 760, 320, 280, 90),
        Box("Capture User Inputs", 760, 460, 280, 90),
        Box("Validate Inputs", 760, 600, 280, 90),
        Box("Generate Diet + Fitness Plan", 690, 740, 420, 100),
        Box("Display Results", 760, 900, 280, 90, fill="#E8F3FF"),
    ]
    for i, s in enumerate(steps):
        draw_box(draw, s)
        if i > 0:
            prev = steps[i - 1]
            arrow(draw, (prev.x + prev.w // 2, prev.y + prev.h), (s.x + s.w // 2, s.y))
    draw.polygon([(1280, 620), (1430, 700), (1280, 780), (1130, 700)], fill="white", outline=BOX_BORDER, width=4)
    draw.text((1182, 688), "Valid?", font=FONT_TEXT, fill=TEXT)
    arrow(draw, (1040, 645), (1130, 700))
    arrow(draw, (1430, 700), (1540, 700))
    draw.text((1550, 687), "No -> show errors", font=FONT_SMALL, fill=TEXT)
    arrow(draw, (1280, 780), (900, 740))
    draw.text((1210, 800), "Yes", font=FONT_SMALL, fill=TEXT)
    save(img, path)


def build_control_flow(path: Path) -> None:
    img, draw = make_canvas("Control Flow Diagram", "Validation and Exception Paths")
    blocks = [
        Box("User Action Event", 150, 260),
        Box("Route Guard", 520, 260),
        Box("Auth Check", 890, 260),
        Box("Input Validator", 1260, 260),
        Box("Recommendation Engine", 700, 520, 400, 120),
        Box("Error Handler", 1260, 520, 300, 120, fill="#FFF1F2"),
        Box("Response Renderer", 700, 780, 400, 120, fill="#E8F3FF"),
    ]
    for b in blocks:
        draw_box(draw, b)
    arrow(draw, (450, 320), (520, 320))
    arrow(draw, (820, 320), (890, 320))
    arrow(draw, (1190, 320), (1260, 320))
    arrow(draw, (1410, 380), (900, 520))
    draw.text((1120, 430), "valid", font=FONT_SMALL, fill=TEXT)
    arrow(draw, (1410, 380), (1410, 520), color="#D62828")
    draw.text((1430, 450), "invalid", font=FONT_SMALL, fill="#D62828")
    arrow(draw, (900, 640), (900, 780))
    arrow(draw, (1260, 580), (1100, 840), color="#D62828")
    save(img, path)


def build_state(path: Path) -> None:
    img, draw = make_canvas("State Transition Diagram", "User Session and Planning States")
    states = [
        ("Idle", (180, 500, 340, 660)),
        ("Authenticated", (460, 320, 730, 500)),
        ("Profile Captured", (800, 320, 1120, 500)),
        ("Plan Generating", (1180, 320, 1510, 500)),
        ("Plan Generated", (900, 640, 1210, 820)),
        ("Displayed", (1290, 640, 1560, 820)),
    ]
    for name, rect in states:
        draw.rounded_rectangle(rect, radius=30, fill=BOX_FILL, outline=BOX_BORDER, width=4)
        tw = draw.textlength(name, font=FONT_LABEL)
        draw.text((rect[0] + (rect[2] - rect[0] - tw) / 2, rect[1] + 70), name, font=FONT_LABEL, fill=TEXT)
    arrow(draw, (340, 580), (460, 410))
    arrow(draw, (730, 410), (800, 410))
    arrow(draw, (1120, 410), (1180, 410))
    arrow(draw, (1345, 500), (1055, 640))
    arrow(draw, (1210, 730), (1290, 730))
    draw.text((580, 560), "login", font=FONT_SMALL, fill=TEXT)
    draw.text((860, 540), "submit profile", font=FONT_SMALL, fill=TEXT)
    draw.text((1200, 560), "compute", font=FONT_SMALL, fill=TEXT)
    draw.text((1160, 690), "success", font=FONT_SMALL, fill=TEXT)
    save(img, path)


def build_architecture(path: Path) -> None:
    img, draw = make_canvas("System Architecture Diagram", "React + FastAPI Modular Architecture")
    ui = Box("Client Browser\nReact + TypeScript", 140, 360, 360, 170, fill="#E8F3FF")
    api = Box("API Gateway\nFastAPI Router", 680, 300, 360, 170)
    diet = Box("Diet Service\nBMI / Macros / Meals", 1180, 220, 420, 170, fill="#ECFDF5")
    fitness = Box("Fitness Service\nWeekly Workouts", 1180, 460, 420, 170, fill="#ECFDF5")
    model = Box("Indian Foods Model\nCurated Meal Dataset", 680, 620, 360, 170)
    for b in [ui, api, diet, fitness, model]:
        draw_box(draw, b)
    arrow(draw, (500, 445), (680, 385))
    arrow(draw, (680, 420), (500, 475))
    arrow(draw, (1040, 385), (1180, 305))
    arrow(draw, (1040, 385), (1180, 545))
    arrow(draw, (860, 620), (860, 470))
    save(img, path)


def build_activity(path: Path) -> None:
    img, draw = make_canvas("Activity Diagram", "NutriSense User Journey")
    steps = [
        Box("Start", 120, 420, 180, 90, fill="#ECFDF5"),
        Box("Open App", 370, 420, 220, 90),
        Box("Login", 660, 420, 180, 90),
        Box("Fill Profile", 900, 420, 240, 90),
        Box("Generate Plan", 1210, 420, 260, 90),
        Box("View Results", 1520, 420, 220, 90, fill="#E8F3FF"),
    ]
    for i, s in enumerate(steps):
        draw_box(draw, s)
        if i > 0:
            p = steps[i - 1]
            arrow(draw, (p.x + p.w, p.y + p.h // 2), (s.x, s.y + s.h // 2))
    save(img, path)


def build_sequence(path: Path) -> None:
    img, draw = make_canvas("Sequence Diagram", "Generate Personalized Plan")
    lifelines = [
        ("User", 220),
        ("Frontend", 560),
        ("API Router", 900),
        ("Diet Service", 1240),
        ("Fitness Service", 1540),
    ]
    for name, x in lifelines:
        draw.rectangle([x - 90, 170, x + 90, 240], fill=BOX_FILL, outline=BOX_BORDER, width=3)
        tw = draw.textlength(name, font=FONT_TEXT)
        draw.text((x - tw / 2, 195), name, font=FONT_TEXT, fill=TEXT)
        draw.line([(x, 240), (x, 980)], fill="#94A3B8", width=3)

    def msg(y: int, x1: int, x2: int, text: str) -> None:
        arrow(draw, (x1, y), (x2, y), color=SECONDARY, width=4)
        draw.text((min(x1, x2) + 20, y - 30), text, font=FONT_SMALL, fill=TEXT)

    msg(320, 220, 560, "Submit profile")
    msg(410, 560, 900, "POST /generate-plan")
    msg(500, 900, 1240, "compute diet")
    msg(590, 900, 1540, "compute workout")
    msg(690, 1240, 900, "diet result")
    msg(770, 1540, 900, "fitness result")
    msg(860, 900, 560, "JSON response")
    msg(940, 560, 220, "Render plan")
    save(img, path)


def build_class(path: Path) -> None:
    img, draw = make_canvas("Class Diagram", "Core Domain and Service Classes")
    classes = [
        Box("UserProfile\n+name\n+age\n+height\n+weight\n+dietType", 120, 240, 320, 280),
        Box("DietService\n+calculateBMI()\n+calculateProtein()\n+generateMealPlan()", 600, 180, 380, 280),
        Box("FitnessService\n+generateWorkoutPlan()", 600, 560, 380, 220),
        Box("IndianFoods\n+foods[]\n+filterByDiet()\n+filterByBudget()", 1140, 180, 430, 260),
        Box("RecommendationResponse\n+mealPlan\n+fitnessPlan\n+bmi\n+goal", 1140, 560, 430, 240),
    ]
    for c in classes:
        draw_box(draw, c)
    arrow(draw, (440, 360), (600, 320), color=ACCENT)
    arrow(draw, (980, 300), (1140, 300), color=ACCENT)
    arrow(draw, (980, 680), (1140, 680), color=ACCENT)
    arrow(draw, (440, 380), (600, 650), color=ACCENT)
    save(img, path)


def build_package(path: Path) -> None:
    img, draw = make_canvas("Package Diagram", "High-Level Package Dependencies")
    pkgs = [
        Box("frontend\n(pages, routes, components)", 140, 320, 380, 200),
        Box("api\n(apiRouter)", 700, 230, 260, 160),
        Box("services\n(dietService, fitnessService)", 1140, 200, 430, 190),
        Box("models\n(indianFoods)", 1140, 510, 320, 170),
        Box("theme/context\n(ui personalization)", 140, 610, 380, 180),
    ]
    for p in pkgs:
        draw_box(draw, p)
    arrow(draw, (520, 420), (700, 300))
    arrow(draw, (960, 300), (1140, 280))
    arrow(draw, (1300, 390), (1300, 510))
    arrow(draw, (520, 700), (700, 330))
    save(img, path)


def build_component(path: Path) -> None:
    img, draw = make_canvas("Component Diagram", "Runtime Components and Interfaces")
    comps = [
        Box("Landing/Login UI", 120, 240, 300, 120),
        Box("User Details UI", 120, 430, 300, 120),
        Box("Plan Result UI", 120, 620, 300, 120),
        Box("Axios API Client", 520, 430, 320, 120),
        Box("FastAPI Router", 950, 430, 300, 120),
        Box("Diet Engine", 1340, 280, 300, 120),
        Box("Fitness Engine", 1340, 570, 300, 120),
    ]
    for c in comps:
        draw_box(draw, c)
    arrow(draw, (420, 300), (520, 470))
    arrow(draw, (420, 490), (520, 490))
    arrow(draw, (420, 680), (520, 530))
    arrow(draw, (840, 490), (950, 490))
    arrow(draw, (1250, 470), (1340, 340))
    arrow(draw, (1250, 510), (1340, 630))
    save(img, path)


def build_deployment(path: Path) -> None:
    img, draw = make_canvas("Deployment Diagram", "Development/Production Node View")
    nodes = [
        Box("Client Device\nWeb Browser", 120, 390, 320, 170, fill="#E8F3FF"),
        Box("Frontend Node\nVite / Static Host", 620, 220, 360, 170),
        Box("Backend Node\nFastAPI Server", 620, 560, 360, 170),
        Box("Data Node (Future)\nRelational DB", 1180, 390, 360, 170, fill="#ECFDF5"),
    ]
    for n in nodes:
        draw_box(draw, n)
    arrow(draw, (440, 475), (620, 305))
    draw.text((470, 360), "HTTPS", font=FONT_SMALL, fill=TEXT)
    arrow(draw, (440, 475), (620, 645))
    draw.text((470, 560), "REST", font=FONT_SMALL, fill=TEXT)
    arrow(draw, (980, 645), (1180, 475))
    draw.text((1020, 540), "ORM / SQL", font=FONT_SMALL, fill=TEXT)
    save(img, path)


def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def insert_images_into_docx(input_docx: Path, output_docx: Path, diagram_dir: Path) -> None:
    doc = Document(str(input_docx))

    targets = [
        ("2.2 Use Case diagram", "use_case.png"),
        ("2.3 ER Model", "er_diagram.png"),
        ("2.4 Data flow Diagram", "dfd.png"),
        ("2.5 Control Flow Diagram", "control_flow.png"),
        ("2.6 State Transition Diagram", "state_transition.png"),
        ("3.1.1 System Architectural Diagram", "system_architecture.png"),
        ("3.3.1 Flowchart", "flowchart.png"),
        ("2.3 Activity Diagram", "activity_diagram.png"),
        ("2.4 Sequence Diagram", "sequence_diagram.png"),
        ("2.5 Class Diagram", "class_diagram.png"),
        ("3.3.1 Package Diagram", "package_diagram.png"),
        ("3.3.2 Component Diagram", "component_diagram.png"),
        ("3.3.2 Deployment Diagram", "deployment_diagram.png"),
    ]

    inserted = set()
    for para in doc.paragraphs:
        text = para.text.strip()
        for heading, image_name in targets:
            if heading in text and heading not in inserted:
                image_path = diagram_dir / image_name
                if image_path.exists():
                    cap_para = insert_paragraph_after(para, f"Figure: {heading}")
                    img_para = insert_paragraph_after(cap_para)
                    run = img_para.add_run()
                    run.add_picture(str(image_path), width=Inches(6.5))
                    inserted.add(heading)
                break

    doc.save(str(output_docx))


def main() -> None:
    project_root = Path(__file__).resolve().parents[1]
    diagram_dir = project_root / "docs" / "sds-diagrams"
    input_docx = project_root / "NutriSense_SDS_Modified.docx"
    output_docx = project_root / "NutriSense_SDS_With_Diagrams.docx"

    builders = [
        ("use_case.png", build_use_case),
        ("er_diagram.png", build_er),
        ("dfd.png", build_dfd),
        ("control_flow.png", build_control_flow),
        ("state_transition.png", build_state),
        ("system_architecture.png", build_architecture),
        ("flowchart.png", build_flow),
        ("activity_diagram.png", build_activity),
        ("sequence_diagram.png", build_sequence),
        ("class_diagram.png", build_class),
        ("package_diagram.png", build_package),
        ("component_diagram.png", build_component),
        ("deployment_diagram.png", build_deployment),
    ]

    for file_name, fn in builders:
        fn(diagram_dir / file_name)

    insert_images_into_docx(input_docx, output_docx, diagram_dir)
    print(f"Generated diagrams in: {diagram_dir}")
    print(f"Created final document: {output_docx}")


if __name__ == "__main__":
    main()
